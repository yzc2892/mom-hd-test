# -*- coding: utf-8 -*-
"""
Convert manuscript.tex -> manuscript.docx
- Equations rendered as crisp images (matplotlib mathtext) and embedded inline / centered.
- The 6 simulation figures embedded inside the Word file.
- Tables (results_table, timing_table, real-data table) rebuilt as native Word tables.
- Citations resolved to (Author, Year) via refs.bib.
- Formatting: Times New Roman 11pt, justified, A4 1in margins, styled headings/theorems/captions.
"""
import os, re, io, textwrap, tempfile
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.image import imread
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
RESULTS = os.path.join(HERE, "results")
ASSETS = tempfile.mkdtemp(prefix="docx_assets_")

# LaTeX accent map (used by text + bibliography cleanup)
_ACCENT = {
    "\\'e": "é", "\\`e": "è", "\\^e": "ê", "\\\"e": "ë",
    "\\'a": "á", "\\`a": "à", "\\^a": "â", "\\\"a": "ä",
    "\\'o": "ó", "\\`o": "ò", "\\^o": "ô", "\\\"o": "ö",
    "\\'u": "ú", "\\`u": "ù", "\\^u": "û", "\\\"u": "ü",
    "\\'i": "í", "\\`i": "ì", "\\^i": "î", "\\'c": "ć",
    "\\c c": "ç", "\\~n": "ñ", "\\'s": "ś", "\\ss": "ß",
    "\\aa": "å", "\\AA": "Å", "\\o": "ø", "\\O": "Ø",
    "\\ae": "æ", "\\AE": "Æ", "\\l": "ł", "\\L": "Ł",
    "\\'E": "É", "\\`E": "È", "\\^E": "Ê",
}

# ----------------------------------------------------------------------------
# 1. Math rendering
# ----------------------------------------------------------------------------
def normalize_math(s):
    s = re.sub(r"\\label\{[^}]*\}", "", s)
    s = re.sub(r"\\normt\{([^{}]*)\}", r"\\left\\lVert \1 \\right\\rVert_2", s)
    s = re.sub(r"\\normi\{([^{}]*)\}", r"\\left\\lVert \1 \\right\\rVert_\\infty", s)
    s = re.sub(r"\\norm\{([^{}]*)\}", r"\\left\\lVert \1 \\right\\rVert", s)
    repl = {
        r"\bX": r"\mathbf{X}", r"\bY": r"\mathbf{Y}",
        r"\bmu": r"\boldsymbol{\mu}", r"\bSig": r"\boldsymbol{\Sigma}",
        r"\bzero": r"\mathbf{0}", r"\R": r"\mathbb{R}", r"\E": r"\mathbb{E}",
        r"\cB": r"\mathcal{B}", r"\cK": r"\mathcal{K}",
        r"\Var": r"\operatorname{Var}", r"\Cov": r"\operatorname{Cov}",
        r"\med": r"\operatorname{med}", r"\tr": r"\operatorname{tr}",
    }
    for k, v in repl.items():
        s = s.replace(k, v)
    s = s.replace(r"\boldsymbol", r"\mathbf")
    s = s.replace(r"\operatorname", r"\mathrm")
    s = re.sub(r"\\mathchar`-`", "-", s)
    s = re.sub(r"\\xrightarrow(\[[^\]]*\])?\{[^}]*\}", r"\\rightarrow", s)
    s = s.replace(r"\widehat", r"\hat")
    s = s.replace(r"\widetilde", r"\tilde")
    return s

_MATH_CACHE = {}

def render_math(tex, display=False):
    key = ("D" if display else "I") + normalize_math(tex).strip()
    if key in _MATH_CACHE:
        return _MATH_CACHE[key]
    clean = normalize_math(tex).strip()
    if not clean:
        return None
    dpi = 200
    fs = 12 if display else 11
    fig = plt.figure(figsize=(8, 0.7))
    try:
        fig.text(0.01, 0.5, f"${clean}$", fontsize=fs, va="center")
        path = os.path.join(ASSETS, f"m_{abs(hash(key)) % 10**9}.png")
        fig.savefig(path, dpi=dpi, bbox_inches="tight", transparent=True,
                    pad_inches=0.02)
        plt.close(fig)
        arr = imread(path)
        h, w = arr.shape[0], arr.shape[1]
        info = (path, w / dpi, h / dpi)  # width_in, height_in
        _MATH_CACHE[key] = info
        return info
    except Exception as e:
        plt.close(fig)
        # fallback: render raw text
        fig = plt.figure(figsize=(8, 0.7))
        fig.text(0.01, 0.5, clean, fontsize=fs, va="center")
        path = os.path.join(ASSETS, f"m_{abs(hash(key)) % 10**9}.png")
        fig.savefig(path, dpi=dpi, bbox_inches="tight", transparent=True,
                    pad_inches=0.02)
        plt.close(fig)
        arr = imread(path)
        h, w = arr.shape[0], arr.shape[1]
        info = (path, w / dpi, h / dpi)
        _MATH_CACHE[key] = info
        return info

# ----------------------------------------------------------------------------
# 2. Bibliography parsing
# ----------------------------------------------------------------------------
def parse_bib(path):
    txt = open(path, encoding="utf-8").read()
    out = {}
    for m in re.finditer(r"@\w+\{([^,]+),\s*(.*?)\n\}", txt, re.DOTALL):
        key = m.group(1).strip()
        body = m.group(2)
        fields = {}
        for fm in re.finditer(r"(\w+)\s*=\s*\{([^{}]*)\}", body, re.DOTALL):
            fields[fm.group(1).lower()] = fm.group(2).strip()
        author = fields.get("author", "")
        year = fields.get("year", "?")
        out[key] = (short_author(author), year)
    return out

def short_author(s):
    s = re.sub(r"\s+", " ", s)
    parts = [p.strip() for p in re.split(r"\s+and\s+", s)]
    parts = [p for p in parts if p and p.lower() != "others"]
    lasts = []
    for p in parts:
        if "," in p:
            lasts.append(p.split(",")[0].strip())
        else:
            sp = p.split()
            lasts.append(sp[-1] if sp else p)
    if not lasts:
        return "Anonymous"
    if len(lasts) == 1:
        return lasts[0]
    if len(lasts) == 2:
        return f"{lasts[0]} and {lasts[1]}"
    return f"{lasts[0]} et al."

BIB = parse_bib(os.path.join(HERE, "refs.bib"))

# citation numbering map (key -> [n]), built by first-appearance order
CITENUM = {}

def build_cite_numbers(text):
    """Assign each cited key a number by order of first appearance
    (mirrors unsrtnat's numerical citation order)."""
    global CITENUM
    CITENUM = {}
    for m in re.finditer(r"\\(?:citep|citeauthor|cite)\{([^}]+)\}", text):
        for k in m.group(1).split(","):
            k = k.strip()
            if k and k not in CITENUM:
                CITENUM[k] = len(CITENUM) + 1

def cite_text(keys, paren):
    """Numerical in-text citation, e.g. [1] or [1, 3]."""
    nums = []
    for k in keys:
        k = k.strip()
        nums.append(str(CITENUM[k]) if k in CITENUM else k)
    return "[" + ", ".join(nums) + "]"

def cite_author_text(keys):
    """natbib-style author names for \\citeauthor (e.g. 'Bai and Saranadasa',
    'Srivastava et al.')."""
    items = []
    for k in keys:
        k = k.strip()
        if k in BIB_FULL:
            field = BIB_FULL[k]["fields"].get("author", "")
            names = bib_author_list(field)
            if not names:
                items.append(k)
            elif len(names) == 1:
                items.append(names[0].split()[-1])
            elif len(names) == 2:
                items.append(names[0].split()[-1] + " and " + names[1].split()[-1])
            else:
                items.append(names[0].split()[-1] + " et al.")
        else:
            items.append(k)
    return "; ".join(items)

# ----------------------------------------------------------------------------
# 2b. Full bibliography parsing for the reference list
# ----------------------------------------------------------------------------
def parse_bib_full(path):
    txt = open(path, encoding="utf-8").read()
    out = {}
    for m in re.finditer(r"@(\w+)\{([^,]+),\s*(.*?)\n\}", txt, re.DOTALL):
        key = m.group(2).strip()
        btype = m.group(1).lower()
        body = m.group(3)
        fields = {}
        for fm in re.finditer(r"(\w+)\s*=\s*\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}", body, re.DOTALL):
            fields.setdefault(fm.group(1).lower(), fm.group(2).strip())
        for fm in re.finditer(r'(\w+)\s*=\s*"([^"]*)"', body):
            fields.setdefault(fm.group(1).lower(), fm.group(2).strip())
        out[key] = {"type": btype, "fields": fields}
    return out

BIB_FULL = parse_bib_full(os.path.join(HERE, "refs.bib"))

def bib_author_list(field):
    """Return list of 'F. M. Last'."""
    if not field:
        return []
    parts = [p.strip() for p in re.split(r"\s+and\s+", field, flags=re.IGNORECASE)]
    parts = [p for p in parts if p and p.lower() != "others"]
    out = []
    for p in parts:
        if not p:
            continue
        if "," in p:
            last, rest = p.split(",", 1)
            first = rest.strip()
        else:
            toks = p.split()
            last = toks[-1]
            first = " ".join(toks[:-1])
        inits = "".join(t[0].upper() + "." for t in first.split() if t)
        out.append(f"{inits} {last}".strip())
    return out

def fmt_authors_for_bib(field, cap=10):
    names = bib_author_list(field)
    if not names:
        return "Anonymous"
    if len(names) <= 2:
        return " and ".join(names)
    if len(names) <= cap:
        return ", ".join(names[:-1]) + ", and " + names[-1]
    return ", ".join(names[:cap]) + ", et al."

def clean_bib_text(s):
    if not s:
        return ""
    s = s.replace("$", "")
    # common math symbols that may appear in titles
    s = s.replace("\\ell", "ℓ")
    s = s.replace("\\alpha", "α").replace("\\beta", "β")
    s = s.replace("\\mu", "μ").replace("\\Sigma", "Σ")
    s = s.replace("\\le", "≤").replace("\\ge", "≥")
    for k, v in _ACCENT.items():
        s = s.replace(k, v)
    s = s.replace("\\&", "&").replace("\\%", "%").replace("\\_", "_")
    s = s.replace("\\#", "#").replace("\\{", "{").replace("\\}", "}")
    s = s.replace("~", " ").replace("---", "—").replace("--", "–")
    # subscripts like _1 -> ₁
    s = re.sub(r"_(\d)", lambda m: "₀₁₂₃₄₅₆₇₈₉"[int(m.group(1))], s)
    s = re.sub(r"\\[a-zA-Z]+\*?", "", s)
    s = s.replace("{", "").replace("}", "")
    s = re.sub(r"\s+", " ", s).strip()
    return s

def bib_sort_key(key, entry):
    names = bib_author_list(entry["fields"].get("author", ""))
    last = names[0].split()[-1].lower() if names else "zzzz"
    year = entry["fields"].get("year", "9999")
    return (last, year, key)

def build_ref_runs(key, entry):
    f = entry["fields"]
    authors = clean_bib_text(fmt_authors_for_bib(f.get("author", "")))
    year = f.get("year", "?")
    title = clean_bib_text(f.get("title", ""))
    btype = entry["type"]
    segs = [(f"{authors}. ", ""), (f"{title}. ", "")]
    if btype == "article":
        journal = clean_bib_text(f.get("journal", ""))
        if journal:
            vol = f.get("volume", "")
            no = f.get("number", "")
            pages = f.get("pages", "").replace("--", "–")
            tail = ""
            if vol:
                tail += f"vol. {vol}, "
            if no:
                tail += f"no. {no}, "
            if pages:
                tail += f"pp. {pages}, "
            tail += f"{year}."
            segs.append((journal, "italic"))
            segs.append((", " + tail, ""))
        else:
            segs.append((f"{year}.", ""))
    elif btype in ("book", "inbook"):
        pub = clean_bib_text(f.get("publisher", ""))
        segs.append((f"{pub}, {year}." if pub else f"{year}.", ""))
    elif btype in ("inproceedings", "conference", "incollection"):
        bt = clean_bib_text(f.get("booktitle", ""))
        pages = f.get("pages", "").replace("--", "–")
        tail = bt if bt else ""
        if pages:
            tail += f", pp. {pages}"
        tail += f", {year}."
        segs.append((tail if tail else f"{year}.", ""))
    elif btype in ("phdthesis", "mastersthesis"):
        school = clean_bib_text(f.get("school", ""))
        segs.append((f"PhD thesis, {school}, {year}." if school else f"{year}.", ""))
    elif btype == "misc":
        note = clean_bib_text(f.get("note", ""))
        how = clean_bib_text(f.get("howpublished", ""))
        if note:
            segs.append((note + " ", ""))
        elif how:
            segs.append((how + ", ", ""))
        segs.append((f"{year}.", ""))
    else:
        segs.append((f"{year}.", ""))
    return segs

def build_references(doc):
    # list only cited entries, in citation order (matches numerical in-text [n])
    entries = [(CITENUM.get(k, 10**9), k, e) for k, e in BIB_FULL.items() if k in CITENUM]
    entries.sort()
    # unnumbered heading
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run("References")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    rPr = r._element.get_or_add_rPr()
    rfonts = rPr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rPr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    for _, key, entry in entries:
        segs = build_ref_runs(key, entry)
        num = CITENUM.get(key)
        if num:
            segs = [("[%d] " % num, "")] + segs
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.left_indent = Cm(0.6)
        para.paragraph_format.first_line_indent = Cm(-0.6)
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.line_spacing = 1.12
        for text, style in segs:
            run = para.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            run.italic = (style == "italic")
            rPr = run._element.get_or_add_rPr()
            rfonts = rPr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rPr.append(rfonts)
            rfonts.set(qn("w:ascii"), "Times New Roman")
            rfonts.set(qn("w:hAnsi"), "Times New Roman")

# ----------------------------------------------------------------------------
# 3. Converter
# ----------------------------------------------------------------------------
STYLE = {
    "Normal": dict(align="justify", size=11, space_after=6, line=1.5),
    "abstract": dict(align="justify", size=10, space_after=4, line=1.25, indent=0.9),
    "theorem_body": dict(align="justify", size=10.5, space_after=4, line=1.35, italic=True),
    "proof_body": dict(align="justify", size=10.5, space_after=4, line=1.35),
    "item": dict(align="justify", size=11, space_after=3, line=1.35, indent=0.6),
    "caption": dict(align="center", size=9.5, italic=True, space_after=8, space_before=2),
    "equation": dict(align="center", size=11, space_after=6, space_before=6),
}

class Converter:
    def __init__(self, doc):
        self.doc = doc
        self.cur = None
        self.cur_style = None
        self.cur_size = 11
        self.cur_italic = False
        self.ctx_stack = ["Normal"]
        self.env_counters = {n: 0 for n in
            ["theorem", "lemma", "corollary", "proposition",
             "assumption", "definition", "remark"]}
        self.sec = 0
        self.subsec = 0
        self.eq = 0
        self.fig = 0
        self.tab = 0
        self.algo = 0
        self.list_depth = 0
        self.labelmap = {}
        self.pending = None  # (kind, number)

    # ---- paragraph helpers ----
    def new_para(self, style, indent=None):
        p = self.doc.add_paragraph()
        spec = STYLE.get(style, STYLE["Normal"])
        al = spec.get("align", "justify")
        p.alignment = {"justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
                       "center": WD_ALIGN_PARAGRAPH.CENTER,
                       "left": WD_ALIGN_PARAGRAPH.LEFT,
                       "right": WD_ALIGN_PARAGRAPH.RIGHT}[al]
        p.paragraph_format.space_after = Pt(spec.get("space_after", 6))
        p.paragraph_format.space_before = Pt(spec.get("space_before", 0))
        p.paragraph_format.line_spacing = spec.get("line", 1.15)
        ind = indent if indent is not None else spec.get("indent")
        if ind is not None:
            p.paragraph_format.left_indent = Cm(ind)
        self.cur = p
        self.cur_style = style
        self.cur_size = spec.get("size", 11)
        self.cur_italic = spec.get("italic", False)
        return p

    def ensure_para(self):
        if self.cur is None or self.cur_style != self.ctx_stack[-1]:
            self.new_para(self.ctx_stack[-1])

    def close_para(self):
        self.cur = None
        self.cur_style = None

    def set_run_font(self, r, size=None, bold=False, italic=False, mono=False):
        r.font.name = "Courier New" if mono else "Times New Roman"
        r.font.size = Pt(size if size else self.cur_size)
        r.bold = bold
        r.italic = italic or self.cur_italic
        rPr = r._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), r.font.name)
        rFonts.set(qn("w:hAnsi"), r.font.name)
        rFonts.set(qn("w:eastAsia"), "SimSun")

    # ---- text emission ----
    def emit_text(self, text, bold=False, italic=False, mono=False, target=None):
        text = replace_cites_refs(text)
        # split on math delimiters $
        parts = text.split("$")
        para = target if target is not None else None
        for i, seg in enumerate(parts):
            if i % 2 == 1:  # math
                if seg.strip():
                    self.add_math_inline(seg.strip(), target=target)
            else:
                seg = clean_text(seg)
                if not seg:
                    continue
                # split into sub-runs on \\ (newline)
                for j, line in enumerate(seg.split("\n")):
                    if j > 0 and (target is not None or self.cur is not None):
                        (target if target is not None else self.cur).add_run().add_break()
                    if not line:
                        continue
                    if target is not None:
                        r = target.add_run(line)
                        self.set_run_font(r, bold=bold, italic=italic, mono=mono)
                    else:
                        self.ensure_para()
                        r = self.cur.add_run(line)
                        self.set_run_font(r, bold=bold, italic=italic, mono=mono)

    def add_math_inline(self, tex, target=None):
        info = render_math(tex, display=False)
        if info is None:
            return
        path, w_in, h_in = info
        height_cm = min(h_in * 2.54, 0.6)
        if target is not None:
            target.add_run().add_picture(path, height=Cm(height_cm))
        else:
            self.ensure_para()
            self.cur.add_run().add_picture(path, height=Cm(height_cm))

    def add_math_display(self, tex, num=None):
        info = render_math(tex, display=True)
        if info is None:
            return
        path, w_in, h_in = info
        self.close_para()
        if num is not None:
            t = self.doc.add_table(rows=1, cols=2)
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            t.autofit = False
            c0 = t.cell(0, 0)
            c1 = t.cell(0, 1)
            c0.width = Cm(13.5)
            c1.width = Cm(2.4)
            c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            pic_w = min(w_in * 2.54, 13.0)
            c0.paragraphs[0].add_run().add_picture(path, width=Cm(pic_w))
            rn = c1.paragraphs[0].add_run(f"({num})")
            self.set_run_font(rn, size=10, bold=False)
            remove_all_borders(t)
        else:
            p = self.new_para("equation")
            max_w = 15.5
            width_cm = min(w_in * 2.54, max_w)
            p.add_run().add_picture(path, width=Cm(width_cm))
        self.close_para()

    # ---- main parser ----
    def parse(self, text):
        text = strip_comments(text)
        PAT = re.compile(r"""
            (?P<env>\\(?:begin|end)\{[a-zA-Z*]+\})
            |(?P<input>\\input\{[^}]*\})
            |(?P<mathdel>\$|\\\[|\\\]|\\\(|\\\))
            |(?P<cmd>\\[a-zA-Z]+\*?|\\[^a-zA-Z%])
        """, re.VERBOSE)
        pos = 0
        n = len(text)
        while pos < n:
            m = None
            for mm in PAT.finditer(text, pos):
                m = mm
                break
            if m is None:
                self.emit_text(text[pos:])
                break
            if m.start() > pos:
                self.emit_text(text[pos:m.start()])
            tok = m.group(0)
            kind = m.lastgroup
            if kind == "mathdel":
                pos = self.handle_mathdel(tok, m.end(), text)
            elif kind == "input":
                pos = self.handle_input(tok, m.end(), text)
            elif kind == "env":
                pos = self.handle_env(tok, m.end(), text)
            else:  # cmd
                pos = self.handle_cmd(tok, m.end(), text)
        self.close_para()

    # ---- handlers ----
    def handle_mathdel(self, tok, end, text):
        if tok == "$":
            close = text.find("$", end)
            if close == -1:
                return len(text)
            self.add_math_inline(text[end:close])
            return close + 1
        if tok == "\\[":
            close = text.find("\\]", end)
            if close == -1:
                return len(text)
            self.add_math_display(text[end:close])
            return close + 2
        if tok == "\\(":
            close = text.find("\\)", end)
            if close == -1:
                return len(text)
            self.add_math_inline(text[end:close])
            return close + 2
        # '\]' or '\\)' stray
        return end

    def handle_input(self, tok, end, text):
        fname = tok[tok.find("{")+1:-1]
        if not os.path.isabs(fname):
            fpath = os.path.join(HERE, fname)
            if not os.path.exists(fpath):
                fpath = os.path.join(RESULTS, os.path.basename(fname))
        else:
            fpath = fname
        if os.path.exists(fpath):
            self.parse(open(fpath, encoding="utf-8").read())
        elif not fpath.endswith(".tex") and os.path.exists(fpath + ".tex"):
            self.parse(open(fpath + ".tex", encoding="utf-8").read())
        return end

    def handle_env(self, tok, end, text):
        name = tok[tok.find("{")+1:].rstrip("*").rstrip("}")
        inner, close = extract_env(text, end)
        if name in ("document",):
            pass
        elif name in ("abstract",):
            self.parse_block(inner, "abstract", numbered=False)
        elif name in self.env_counters:
            self.env_counters[name] += 1
            num = self.env_counters[name]
            prefix = getattr(self, "env_prefix", {}).get(name, "")
            self.pending = (name, f"{prefix}{num}")
            title = ""
            optm = re.match(r"\s*\[([^\]]*)\]", inner)
            if optm:
                title = optm.group(1)
                inner = inner[optm.end():]
            self.close_para()
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            cap = name.capitalize()
            r = p.add_run(f"{cap} {prefix}{num}.")
            self.set_run_font(r, size=11, bold=True)
            if title:
                r2 = p.add_run(" " + title)
                self.set_run_font(r2, size=11, bold=True)
            body_style = "theorem_body"
            self.ctx_stack.append(body_style)
            self.parse(inner)
            self.ctx_stack.pop()
            self.close_para()
        elif name == "remark":
            self.env_counters["remark"] += 1
            num = self.env_counters["remark"]
            self.pending = ("remark", num)
            self.close_para()
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(f"Remark {num}.")
            self.set_run_font(r, size=11, bold=True)
            self.ctx_stack.append("proof_body")
            self.parse(inner)
            self.ctx_stack.pop()
            self.close_para()
        elif name == "proof":
            optm = re.match(r"\s*\[([^\]]*)\]", inner)
            head = optm.group(1) if optm else "Proof"
            inner2 = inner[optm.end():] if optm else inner
            self.pending = None
            self.close_para()
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(head + ".")
            self.set_run_font(r, size=11, bold=True)
            self.ctx_stack.append("proof_body")
            self.parse(inner2)
            self.ctx_stack.pop()
            self.close_para()
        elif name == "equation":
            lab = re.search(r"\\label\{[^}]*\}", inner)
            num = self.eq + 1
            self.eq += 1
            if lab:
                self.pending = ("equation", num)
                self.labelmap[lab.group(0)[7:-1]] = ("equation", num)
            self.add_math_display(inner, num)
        elif name == "figure":
            self.handle_figure(inner)
        elif name == "table":
            self.handle_table(inner)
        elif name in ("itemize", "enumerate"):
            self.handle_list(name, inner)
        elif name == "algorithm":
            self.handle_algorithm(inner)
        elif name in ("algorithmic",):
            self.parse(inner)
        elif name in ("tabular",):
            self.build_table(inner, None)
        elif name in ("center", "centerling", "centering"):
            self.parse(inner)
        else:
            # unknown env: just parse inner as text
            self.parse(inner)
        # skip optional args already consumed inside handlers
        return close

    def parse_block(self, inner, style, numbered):
        self.close_para()
        self.ctx_stack.append(style)
        self.parse(inner)
        self.ctx_stack.pop()
        self.close_para()

    def handle_figure(self, inner):
        # find includegraphics
        inc = re.search(r"\\includegraphics(\[[^\]]*\])?\{([^}]*)\}", inner)
        path = None
        width_frac = 0.8
        if inc:
            opt = inc.group(1)
            if opt:
                wm = re.search(r"width=([0-9.]+)\\textwidth", opt)
                if wm:
                    width_frac = float(wm.group(1))
            rel = inc.group(2)
            if os.path.exists(rel):
                path = rel
            elif os.path.exists(os.path.join(HERE, rel)):
                path = os.path.join(HERE, rel)
            elif os.path.exists(os.path.join(RESULTS, os.path.basename(rel))):
                path = os.path.join(RESULTS, os.path.basename(rel))
        cap = find_caption(inner)
        lab = find_label(inner)
        self.fig += 1
        self.pending = ("figure", self.fig)
        if lab:
            self.labelmap[lab] = ("figure", self.fig)
        self.close_para()
        if path:
            p = self.new_para("equation")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            textw = 15.92
            p.add_run().add_picture(path, width=Cm(width_frac * textw))
        if cap:
            cp = self.new_para("caption")
            cp.add_run(f"Figure {self.fig}. ")
            self.emit_text(cap, target=cp)
        self.close_para()

    def handle_table(self, inner):
        cap = find_caption(inner)
        lab = find_label(inner)
        self.tab += 1
        self.pending = ("table", self.tab)
        if lab:
            self.labelmap[lab] = ("table", self.tab)
        # capture caption + label, then render tabulars
        self.close_para()
        caption_para = None
        if cap:
            caption_para = self.new_para("caption")
            caption_para.add_run(f"Table {self.tab}. ")
            self.emit_text(cap, target=caption_para)
        # parse inner for tabulars (and ignore \centering/\small)
        self.parse_tabular_only(inner)
        self.close_para()

    def parse_tabular_only(self, text):
        # find each \begin{tabular}...\end{tabular}
        pos = 0
        while True:
            m = re.search(r"\\begin\{tabular\}", text[pos:])
            if not m:
                break
            begin_abs = pos + m.start()
            inner, close = extract_env(text, pos + m.end())
            # strip the leading column-spec argument {lccc...}
            inner = re.sub(r"^\s*\{[^{}]*\}", "", inner, count=1)
            self.build_table(inner, close)
            pos = close

    def build_table(self, inner, close_hint):
        # remove booktabs rule commands so they don't merge with header rows
        inner = re.sub(r"\\(?:toprule|midrule|bottomrule|hline)\b", " ", inner)
        inner = re.sub(r"\\cmidrule\{[^}]*\}", " ", inner)
        rows = []
        for line in inner.split("\\\\"):
            line = line.strip()
            if not line:
                continue
            if re.match(r"\\(toprule|midrule|bottomrule|cmidrule|hline)\b", line):
                continue
            cells = []
            for c in split_row(line):
                c = c.strip()
                cm = re.match(r"\\multicolumn\{(\d+)\}\{[^}]*\}\{(.*)\}", c, re.DOTALL)
                if cm:
                    cells.append((cm.group(2), int(cm.group(1))))
                else:
                    cells.append((c, 1))
            rows.append(cells)
        if not rows:
            return
        ncols = max(sum(span for _, span in r) for r in rows)
        t = self.doc.add_table(rows=len(rows), cols=ncols)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.style = "Table Grid"
        t.autofit = False
        for i, row in enumerate(rows):
            col = 0
            for celltext, span in row:
                cell = t.cell(i, col)
                if span > 1:
                    merge_to = t.cell(i, min(col + span - 1, ncols - 1))
                    if merge_to != cell:
                        cell.merge(merge_to)
                # clear default paragraph, add our content
                cp = cell.paragraphs[0]
                cp.text = ""
                self.emit_text(celltext, target=cp)
                for r in cp.runs:
                    r.font.size = Pt(9)
                    r.font.name = "Times New Roman"
                if i == 0:
                    for r in cp.runs:
                        r.bold = True
                col += span
        apply_booktabs(t)
        # small spacing after table
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)
        self.close_para()

    def handle_list(self, kind, inner):
        self.close_para()
        self.list_depth += 1
        self.ctx_stack.append("item")
        # parse items
        self.parse_list_items(kind, inner)
        self.ctx_stack.pop()
        self.list_depth -= 1
        self.close_para()

    def parse_list_items(self, kind, inner):
        # split on \item
        parts = re.split(r"\\item\b", inner)
        idx = 0
        for k, part in enumerate(parts):
            if k == 0 and part.strip() == "":
                continue
            if k == 0:
                # content before first \item (rare) - emit as normal
                self.parse(part)
                continue
            self.close_para()
            p = self.new_para("item")
            p.paragraph_format.left_indent = Cm(0.6 + 0.4 * (self.list_depth - 1))
            if kind == "enumerate":
                idx += 1
                r = p.add_run(f"{idx}. ")
                self.set_run_font(r, bold=True)
            else:
                r = p.add_run("• ")
                self.set_run_font(r, bold=True)
            self.parse(part)
        self.close_para()

    def handle_algorithm(self, inner):
        self.algo += 1
        cap = re.search(r"\\caption\{([^{}]*)\}", inner)
        self.close_para()
        if cap:
            cp = self.new_para("caption")
            cp.add_run(f"Algorithm {self.algo}. ")
            self.emit_text(cap.group(1), target=cp)
        # parse algorithmic part
        am = re.search(r"\\begin\{algorithmic\}.*?\\end\{algorithmic\}", inner, re.DOTALL)
        if am:
            alg_inner = am.group(0)
            alg_inner = alg_inner[len("\\begin{algorithmic}"):]
            alg_inner = alg_inner[:alg_inner.rfind("\\end{algorithmic}")]
            self.parse_algorithmic(alg_inner)
        self.close_para()

    def parse_algorithmic(self, inner):
        parts = re.split(r"\\STATE\b", inner)
        idx = 0
        self.ctx_stack.append("item")
        for k, part in enumerate(parts):
            if k == 0:
                if part.strip():
                    self.parse(part)
                continue
            self.close_para()
            p = self.new_para("item")
            p.paragraph_format.left_indent = Cm(0.8)
            idx += 1
            r = p.add_run(f"{idx}. ")
            self.set_run_font(r, bold=True)
            self.parse(part)
        self.close_para()
        self.ctx_stack.pop()

    def handle_cmd(self, tok, end, text):
        name = tok.lstrip("\\").rstrip("*")
        # commands with mandatory braced arg
        if name in ("section", "subsection", "paragraph", "textbf", "textit",
                    "texttt", "text", "emph", "label", "ref", "eqref", "cite",
                    "citep", "citet", "citeauthor", "caption", "thanks", "url", "href",
                    "mathcal", "mathrm", "mathbf", "mathbb"):
            if name in ("citep", "citet", "cite", "citeauthor"):
                # skip optional arguments [..] then read braced key list
                p = end
                while text[p:p+1] == "[":
                    _, p = read_opt(text, p)
                arg, nend = read_braced(text, p)
            else:
                if not text[end:].startswith("{"):
                    import sys
                    sys.stderr.write(f"NOBRACE cmd={name!r} at {end} ctx={text[end-10:end+25]!r}\n")
                    return end
                arg, nend = read_braced(text, end)
            if name == "section":
                self.sec += 1
                self.subsec = 0
                self.pending = ("section", str(self.sec))
                self.close_para()
                self.add_heading(f"{self.sec}  {arg}", 14)
            elif name == "subsection":
                self.subsec += 1
                self.pending = ("subsection", f"{self.sec}.{self.subsec}")
                self.close_para()
                self.add_heading(f"{self.sec}.{self.subsec}  {arg}", 12)
            elif name == "paragraph":
                self.ensure_para()
                r = self.cur.add_run(arg + " ")
                self.set_run_font(r, bold=True)
            elif name in ("textbf", "text", "emph", "textit"):
                style = self.ctx_stack[-1]
                self.ensure_para()
                self.emit_text(arg, bold=(name in ("textbf", "text")),
                               italic=(name in ("textit", "emph")))
            elif name == "texttt":
                self.ensure_para()
                self.emit_text(arg, mono=True)
            elif name in ("citep", "citet", "cite"):
                keys = [k for k in arg.split(",")]
                paren = (name == "citep")
                s = cite_text(keys, paren)
                self.ensure_para()
                self.emit_text(s)
            elif name == "citeauthor":
                keys = [k for k in arg.split(",")]
                s = cite_author_text(keys)
                self.ensure_para()
                self.emit_text(s)
            elif name == "ref":
                self.ensure_para()
                self.emit_text(format_ref(arg))
            elif name == "eqref":
                self.ensure_para()
                self.emit_text(format_ref(arg, eq=True))
            elif name == "label":
                if self.pending:
                    self.labelmap[arg] = self.pending
            elif name == "caption":
                pass  # handled in figure/table
            elif name == "thanks":
                pass
            elif name == "url":
                self.ensure_para()
                self.emit_text(arg, mono=True)
            elif name == "href":
                # \href{url}{text}
                arg2, nend2 = read_braced(text, nend)
                self.ensure_para()
                self.emit_text(arg2)
                nend = nend2
            elif name in ("mathcal", "mathrm", "mathbf", "mathbb"):
                # leftover single-token math command in text -> drop
                pass
            return nend if name != "href" else nend
        if name == "S":
            self.ensure_para()
            self.emit_text("§")
            return end
        if name in ("noindent", "quad", "qquad", ":", ";", ",", "!", " ",
                    "indent", "newline", "par", "medskip", "bigskip", "smallskip"):
            return end
        if name in ("item",):
            return end
        if name == "xrightarrow":
            return end  # should not appear outside math
        # unknown command: skip backslash, keep following chars
        return end

    def add_heading(self, text, size):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12 if size >= 14 else 8)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.bold = True
        rPr = r._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts"); rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")
        self.cur = None
        self.cur_style = None

# ----------------------------------------------------------------------------
# helpers (module level)
# ----------------------------------------------------------------------------
def strip_comments(s):
    out = []
    for line in s.split("\n"):
        # remove % comments not preceded by backslash
        line = re.sub(r"(?<!\\)%.*$", "", line)
        out.append(line)
    return "\n".join(out)

def clean_text(s):
    s = s.replace("\\&", "&").replace("\\%", "%").replace("\\_", "_")
    s = s.replace("\\{", "{").replace("\\}", "}").replace("\\$", "$")
    s = s.replace("\\~", " ").replace("\\-", "-")
    s = s.replace("~", " ")  # LaTeX non-breaking space
    s = s.replace("\\ ", " ").replace("\\\\", "\n")
    s = s.replace("\\S", "§")
    s = s.replace("\\textbackslash", "\\")
    # leftover lone backslash commands -> drop backslash
    s = re.sub(r"\\[a-zA-Z]+\*?", "", s)
    return s

def replace_cites_refs(s):
    s = re.sub(r"\\citep\{([^}]*)\}", lambda m: cite_text(m.group(1).split(","), True), s)
    s = re.sub(r"\\citet\{([^}]*)\}", lambda m: cite_text(m.group(1).split(","), False), s)
    s = re.sub(r"\\cite\{([^}]*)\}", lambda m: cite_text(m.group(1).split(","), True), s)
    s = re.sub(r"\\eqref\{([^}]*)\}", lambda m: format_ref(m.group(1), eq=True), s)
    s = re.sub(r"\\ref\{([^}]*)\}", lambda m: format_ref(m.group(1)), s)
    return s

# ref formatting needs converter's labelmap; set later
_CONV = None
def format_ref(key, eq=False):
    if _CONV is None or key not in _CONV.labelmap:
        return key
    kind, num = _CONV.labelmap[key]
    if kind == "equation":
        return f"({num})"
    return str(num)

def read_opt(text, pos):
    # text[pos] == '['
    assert text[pos] == "["
    depth = 0
    i = pos
    while i < len(text):
        if text[i] == "[":
            depth += 1
        elif text[i] == "]":
            depth -= 1
            if depth == 0:
                return text[pos+1:i], i+1
        i += 1
    return text[pos+1:], len(text)

def read_braced(text, pos):
    assert text[pos] == "{"
    depth = 0
    i = pos
    while i < len(text):
        if text[i] == "{":
            depth += 1
        elif text[i] == "}":
            depth -= 1
            if depth == 0:
                return text[pos+1:i], i+1
        i += 1
    return text[pos+1:], len(text)

def extract_env(text, pos):
    # pos points right after the top-level '\begin{NAME}'
    start = pos
    depth = 1
    i = pos
    n = len(text)
    while i < n and depth > 0:
        if text.startswith("\\begin{", i):
            m = re.match(r"\\begin\{([a-zA-Z*]+)\}", text[i:])
            if m:
                depth += 1
                i += len(m.group(0))
                continue
        if text.startswith("\\end{", i):
            m = re.match(r"\\end\{([a-zA-Z*]+)\}", text[i:])
            if m:
                depth -= 1
                if depth == 0:
                    return text[start:i], i + len(m.group(0))
                i += len(m.group(0))
                continue
        i += 1
    return text[start:], n

def find_caption(inner):
    """Extract caption content allowing nested braces."""
    m = re.search(r"\\caption", inner)
    if not m:
        return None
    bp = inner.find("{", m.end())
    if bp == -1:
        return None
    content, _ = read_braced(inner, bp)
    return content


def find_label(inner):
    m = re.search(r"\\label\{([^}]*)\}", inner)
    return m.group(1) if m else None


def split_row(line):
    # split a tabular row by & respecting \multicolumn braces
    out = []
    depth = 0
    cur = ""
    i = 0
    while i < len(line):
        c = line[i]
        if c == "{":
            depth += 1
            cur += c
        elif c == "}":
            depth -= 1
            cur += c
        elif c == "&" and depth == 0:
            out.append(cur)
            cur = ""
        else:
            cur += c
        i += 1
    out.append(cur)
    return out

def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)

def set_cell_border(cell, edge, val="single", sz="4", color="000000"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    e = borders.find(qn(f"w:{edge}"))
    if e is None:
        e = OxmlElement(f"w:{edge}")
        borders.append(e)
    e.set(qn("w:val"), val)
    e.set(qn("w:sz"), sz)
    e.set(qn("w:space"), "0")
    e.set(qn("w:color"), color)

def apply_booktabs(t):
    """Strip all borders, then add top + mid rules (header) and bottom rule."""
    tbl = t._tbl
    tblPr = tbl.tblPr
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = borders.find(qn(f"w:{edge}"))
        if e is None:
            e = OxmlElement(f"w:{edge}")
            borders.append(e)
        e.set(qn("w:val"), "none")
        e.set(qn("w:sz"), "0")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "auto")
    if len(t.rows) == 0:
        return
    for c in t.rows[0].cells:
        set_cell_border(c, "top")
        set_cell_border(c, "bottom")
    for c in t.rows[-1].cells:
        set_cell_border(c, "bottom")

def remove_all_borders(t):
    tbl = t._tbl
    tblPr = tbl.tblPr
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = borders.find(qn(f"w:{edge}"))
        if e is None:
            e = OxmlElement(f"w:{edge}")
            borders.append(e)
        e.set(qn("w:val"), "none")
        e.set(qn("w:sz"), "0")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "auto")

def add_page_numbers(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._element.append(fld1)
    run._element.append(instr)
    run._element.append(fld2)

def _set_run_ref(r, key):
    kind, num = _CONV.labelmap[key]
    r.text = f"({num})" if kind == "equation" else str(num)

def fix_forward_refs(doc):
    """Replace any unresolved \\ref/\\eqref keys (forward references) with numbers."""
    for p in doc.paragraphs:
        for r in p.runs:
            if r.text in _CONV.labelmap:
                _set_run_ref(r, r.text)
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    for r in p.runs:
                        if r.text in _CONV.labelmap:
                            _set_run_ref(r, r.text)

# ----------------------------------------------------------------------------
# 4. Build document
# ----------------------------------------------------------------------------
def main():
    src = open(os.path.join(HERE, "manuscript.tex"), encoding="utf-8").read()
    build_cite_numbers(src)  # populate CITENUM for numerical in-text citations
    # title & author from preamble
    title_m = re.search(r"\\title\{([^{}]*)\}", src, re.DOTALL)
    author_m = re.search(r"\\author\{(.*?)\n\\date", src, re.DOTALL)
    title = ""
    if title_m:
        title = clean_text(title_m.group(1))
        title = re.sub(r"\s+", " ", title).strip()
    author = ""
    if author_m:
        author = author_m.group(1)
        # remove \thanks{...} with one level of nested braces (e.g. \texttt{...})
        author = re.sub(r"\\thanks\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}", "", author)
        author = author.strip().strip("}").strip("{").strip()
        author = author.replace("\\and", "\n").replace("\\\\", "\n")
        author = clean_text(author)

    # body between \begin{document} and \end{document}
    dm = re.search(r"\\begin\{document\}(.*)\\end\{document\}", src, re.DOTALL)
    body = dm.group(1) if dm else src
    # drop \maketitle, \graphicspath, \bibliography, \bibliographystyle
    body = re.sub(r"\\maketitle", "", body)
    body = re.sub(r"\\graphicspath\{[^}]*\}", "", body)
    body = re.sub(r"\\bibliographystyle\{[^}]*\}", "", body)
    body = re.sub(r"\\bibliography\{[^}]*\}", "", body)
    # drop the appendix "Supplementary Material" note? keep it but it's fine.
    # extract abstract block separately (between \begin{abstract} and \end{abstract})
    abs_m = re.search(r"\\begin\{abstract\}(.*?)\\end\{abstract\}", body, re.DOTALL)
    abstract = abs_m.group(1) if abs_m else ""
    if abs_m:
        body = body[:abs_m.start()] + body[abs_m.end():]

    doc = Document()
    # page setup A4, 1in margins
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    for m in (sec.left_margin, sec.right_margin, sec.top_margin, sec.bottom_margin):
        pass
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    # default font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "SimSun")

    conv = Converter(doc)
    global _CONV
    _CONV = conv

    # Title
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run(title)
    tr.bold = True; tr.font.size = Pt(16); tr.font.name = "Times New Roman"
    if author:
        ap = doc.add_paragraph()
        ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for j, line in enumerate(author.split("\n")):
            if j > 0:
                ap.add_run().add_break()
            ar = ap.add_run(line.strip())
            ar.font.size = Pt(12); ar.font.name = "Times New Roman"

    # Abstract
    if abstract:
        conv.close_para()
        ap = doc.add_paragraph()
        ap.paragraph_format.left_indent = Cm(0.9)
        ap.paragraph_format.right_indent = Cm(0.9)
        ap.paragraph_format.space_after = Pt(8)
        ar = ap.add_run("Abstract")
        ar.bold = True; ar.font.size = Pt(11); ar.font.name = "Times New Roman"
        conv.ctx_stack.append("abstract")
        # parse the abstract body so \textbf / \citet etc. are handled;
        # split on \noindent to keep the structured sub-sections as paragraphs
        for chunk in re.split(r"\\noindent", abstract):
            chunk = chunk.strip()
            if not chunk:
                continue
            conv.close_para()
            p = conv.new_para("abstract")
            p.paragraph_format.right_indent = Cm(0.9)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.25
            conv.cur = p
            conv.cur_style = "abstract"
            conv.cur_size = 10
            conv.cur_italic = False
            conv.parse(chunk)
        conv.ctx_stack.pop()
        conv.close_para()

    # parse body
    conv.parse(body)

    # resolve forward references left as raw keys
    fix_forward_refs(doc)

    # references section
    build_references(doc)

    # page numbers in footer
    add_page_numbers(doc)

    out = os.path.join(HERE, "manuscript.docx")
    doc.save(out)
    print("SAVED", out, "refs:", len(BIB_FULL))

if __name__ == "__main__":
    main()
