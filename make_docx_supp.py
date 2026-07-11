# -*- coding: utf-8 -*-
"""
Convert supplementary.tex -> supplementary.docx
Reuses the converter machinery in make_docx.py (math-as-image, citations,
page numbers, reference list) and adds supplementary-specific formatting:
  - independent title "Supplementary Material to ..."
  - Lemma counters rendered as "Lemma A.1", "Lemma A.2" (via env_prefix)
  - eqref to main-text equations (eq:null / eq:snr / eq:powerrate) resolve to
    their real numbers in the main manuscript (11 / 12 / 13)
"""
import os, re
import importlib.util

# load make_docx as a module (it only runs main() under __main__)
_spec = importlib.util.spec_from_file_location("md", os.path.join(os.path.dirname(os.path.abspath(__file__)), "make_docx.py"))
md = importlib.util.module_from_spec(_spec)
_spec.loader.exec_module(md)

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# main-text equation numbers (verified by order of \begin{equation} in manuscript.tex)
MAIN_EQ_ALIASES = {
    "eq:null": ("equation", 11),
    "eq:snr": ("equation", 12),
    "eq:powerrate": ("equation", 13),
}


class SuppConverter(md.Converter):
    def __init__(self, doc):
        super().__init__(doc)
        self.env_prefix = {"lemma": "A."}
        self.labelmap.update(MAIN_EQ_ALIASES)


def main():
    src = open(os.path.join(md.HERE, "supplementary.tex"), encoding="utf-8").read()

    # ---- title & author ----
    ti = src.find("\\title{")
    title, _ = md.read_braced(src, ti + 6)
    title = md.clean_text(title)
    title = title.replace("``", "“").replace("''", "”")
    title = re.sub(r"\s+", " ", title).strip()

    ai = src.find("\\author{")
    author, _ = md.read_braced(src, ai + 7)
    author = re.sub(r"\\thanks\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}", "", author)
    author = author.replace("\\and", "\n").replace("\\\\", "\n")
    author = md.clean_text(author).strip()

    # ---- body ----
    dm = re.search(r"\\begin\{document\}(.*)\\end\{document\}", src, re.DOTALL)
    body = dm.group(1) if dm else src
    body = re.sub(r"\\maketitle", "", body)
    body = re.sub(r"\\graphicspath\{[^}]*\}", "", body)
    body = re.sub(r"\\bibliographystyle\{[^}]*\}", "", body)
    body = re.sub(r"\\bibliography\{[^}]*\}", "", body)

    # ---- document setup (same as make_docx.main) ----
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement_rfonts()
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "SimSun")

    conv = SuppConverter(doc)
    md._CONV = conv

    # ---- title ----
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run(title)
    tr.bold = True
    tr.font.size = Pt(16)
    tr.font.name = "Times New Roman"
    if author:
        ap = doc.add_paragraph()
        ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for j, line in enumerate(author.split("\n")):
            if j > 0:
                ap.add_run().add_break()
            ar = ap.add_run(line.strip())
            ar.font.size = Pt(12)
            ar.font.name = "Times New Roman"

    # ---- parse body ----
    conv.parse(body)
    md.fix_forward_refs(doc)
    md.build_references(doc)
    md.add_page_numbers(doc)

    out = os.path.join(md.HERE, "supplementary.docx")
    doc.save(out)
    print("SAVED", out)


def OxmlElement_rfonts():
    from docx.oxml import OxmlElement
    return OxmlElement("w:rFonts")


if __name__ == "__main__":
    main()
